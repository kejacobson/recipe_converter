#!/usr/bin/env python
import cv2
import pytesseract
import docx
import re
import os
from pdf2image import convert_from_path
from contextlib import contextmanager

@contextmanager
def cd(path):
    old_dir = os.getcwd()
    os.chdir(path)
    yield
    os.chdir(old_dir)

class RecipeConverter:
    def __init__(self):
        self.folder = 'recipes_to_convert'
        self.word_folder = 'converted_recipes'

    def run(self):
        dir_path = os.path.dirname(os.path.realpath(__file__))

        with cd(dir_path):
            self.remove_ds_store()
            files = os.listdir(self.folder)

            for filename in files:
                image_file = f'{self.folder}/{filename}'
                word_file = f'{self.word_folder}/{self.get_word_file_name(filename)}'

                doc = docx.Document()
                images, image_files = self.read_images_from_file(image_file)
                for image in images:
                    text = self.read_text_from_image(image)
                    self.write_text_to_word_doc(doc, text)

                for image_file in image_files:
                    self.write_image_to_word_doc(doc, image_file)
                doc.save(word_file)

    def read_images_from_file(self, image_file):
        print('Reading', image_file)
        if self.filetype_is_pdf(image_file):
            return self.convert_pdf_to_image(image_file)
        else:
            return [cv2.imread(image_file)], [image_file]

    def filetype_is_pdf(self, image_file):
        return '.pdf' == image_file[-4:]

    def remove_ds_store(self):
        ds_store = f'{self.folder}/.DS_Store'
        if os.path.exists(ds_store):
            os.system(f'rm {ds_store}')

    def get_file_root(self, file):
        return os.path.splitext(file)[0]

    def get_word_file_name(self, image_file):
        return self.get_file_root(image_file) + '.docx'

    def convert_pdf_name_to_image(self, image_filename):
        return self.get_file_root(image_filename)+'.jpg'

    def make_string_xml_compatible(self, line):
        re.sub(u'[^\u0020-\uD7FF\u0009\u000A\u000D\uE000-\uFFFD\U00010000-\U0010FFFF]+', '', line)

    def convert_pdf_to_image(self, image_filename):
        imgs = []
        image_files = []
        images = convert_from_path(image_filename)
        for i, image in enumerate(images):
            jpg_file = self.get_file_root(image_filename) +f'{i}.jpg'
            images[0].save(jpg_file, 'JPEG')
            imgs.append(cv2.imread(jpg_file))
            image_files.append(jpg_file)
        return imgs, image_files

    def read_text_from_image(self, image):
        text = pytesseract.image_to_string(image)
        re.sub('\n\n','\n',text)
        return text

    def write_text_to_word_doc(self, doc, text):
        for line in text.splitlines():
            self.make_string_xml_compatible(line)
            doc.add_paragraph(line)

    def write_image_to_word_doc(self, doc, image_path):
        doc.add_picture(image_path, width=docx.shared.Inches(7))

if __name__ == '__main__':
    converter = RecipeConverter()
    converter.run()
