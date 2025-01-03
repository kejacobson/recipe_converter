#!/usr/bin/env python3
import re
import os
import time
from contextlib import contextmanager
from typing import List
from PIL import Image
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options

import cv2
import pytesseract
import docx
import pdf2image


@contextmanager
def cd(path):
    """
    A context manager to move to a certain directory to execute
    code that will read/write files
    """
    old_dir = os.getcwd()
    os.chdir(path)
    yield
    os.chdir(old_dir)


class RecipeConverter:
    def __init__(self):
        """
        File type converter for recipes.
        Converts recipes from screenshots of websites or pdfs from scans into word documents.
        Uses pytesseract to try to read words from the pictures.
        Includes the original pictures at the bottom of the word document.

        Usage:
        1. Put images and pdfs of recipes into the `recipes_to_convert` directory
        2. Do the `run` command
        """
        self.converter_workspace_dir = "~/Desktop/recipe_converter"
        self.input_folder = "recipes_to_convert"
        self.word_folder = "converted_recipes"

        self.valid_image_types = [".pdf", ".jpg", ".jpeg", ".png", ".jpe", ".bmp", ".jp2", ".tiff", ".tif"]
        self.known_extra_files = [".DS_Store", ".gitkeep"]

    def run(self):
        """
        Finds all the image and pdf files in the `recipes_to_convert` directory
        and parses the strings out of the image files, then writes
        docx files with the strings and images into the `converted_recipes` directory
        """
        with cd(os.path.expanduser(self.converter_workspace_dir)):
            self._make_directory(self.word_folder)
            files_to_convert = os.listdir(self.input_folder)

            for filename in files_to_convert:
                if self._is_a_valid_pdf_or_image_type(filename):
                    self._convert_image_to_word(filename)
                else:
                    if filename not in self.known_extra_files:
                        print(f"Warning unable to convert {filename}. Unknown image extension")

    def _convert_image_to_word(self, filename: str):
        original_image_file = f"{self.input_folder}/{filename}"
        word_file = f"{self.word_folder}/{self._make_word_file_name(filename)}"

        doc = docx.Document()
        images, image_files = self._read_images_from_file(original_image_file)
        self._write_text_section_of_word_doc(doc, images)

        for image_file in image_files:
            self._write_image_to_word_doc(doc, image_file)
            if self._filetype_is_pdf(original_image_file):
                os.system(f'rm "{image_file}"')
        doc.save(word_file)

    def _write_text_section_of_word_doc(self, doc: docx.document.Document, images: List[cv2.Mat]):
        for image in images:
            text = self._read_text_from_image(image)
            self._write_parsed_text_to_word_doc(doc, text)

    def _is_a_valid_pdf_or_image_type(self, filename):
        file_extension = os.path.splitext(filename)[-1]
        return file_extension.lower() in self.valid_image_types

    def _read_text_from_image(self, image):
        image_rgb = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)  # Convert BGR to RGB
        image = Image.fromarray(image_rgb)  # Convert to PIL Image
        max_size = (2000, 2000)
        if image.size[0] > max_size[0] or image.size[1] > max_size[1]:
            image.thumbnail(max_size)
        text = pytesseract.image_to_string(image)
        return text

    def _write_parsed_text_to_word_doc(self, doc: docx.document.Document, text: str):
        for line in text.splitlines():
            self._make_string_xml_compatible(line)
            doc.add_paragraph(line)

    def _make_directory(self, directory: str):
        if not os.path.exists(directory):
            os.mkdir(directory)

    def _read_images_from_file(self, image_filename: str):
        print("Reading", image_filename)
        if self._filetype_is_pdf(image_filename):
            return self._convert_pdf_to_images(image_filename)
        else:
            return [cv2.imread(image_filename)], [image_filename]

    def _filetype_is_pdf(self, image_file: str):
        return ".pdf" == self._get_file_extension(image_file)

    def _get_file_rootname(self, file: str) -> str:
        return os.path.splitext(file)[0]

    def _get_file_extension(self, file: str) -> str:
        return os.path.splitext(file)[-1]

    def _make_word_file_name(self, image_file: str) -> str:
        return self._get_file_rootname(image_file) + ".docx"

    def _convert_pdf_name_to_jpg_image_name(self, image_filename: str) -> str:
        return self._get_file_rootname(image_filename) + ".jpg"

    def _make_string_xml_compatible(self, line: str):
        re.sub("[^\u0020-\uD7FF\u0009\u000A\u000D\uE000-\uFFFD\U00010000-\U0010FFFF]+", "", line)

    def _convert_pdf_to_images(self, pdf_filename: str):
        cv_imgs: List[cv2.Mat] = []
        image_filenames: List[str] = []
        pil_images: List[Image] = pdf2image.convert_from_path(pdf_filename)

        for i, pil_image in enumerate(pil_images):
            file_root = self._get_file_rootname(pdf_filename)
            jpg_file = f"{file_root}{i}.jpg"
            pil_image.save(jpg_file, "JPEG")
            cv_imgs.append(cv2.imread(jpg_file))
            image_filenames.append(jpg_file)
        return cv_imgs, image_filenames

    def _write_image_to_word_doc(self, doc: docx.document.Document, image_path: str):
        doc.add_picture(image_path, width=docx.shared.Inches(7))

    def convert_a_website(self, recipe_name, web_address):
        chrome_options = Options()
        chrome_options.add_argument("--headless")
        chrome_options.add_argument("--disable-gpu")
        chrome_options.add_argument("--no-sandbox")
        service = Service("/usr/local/bin/chromedriver")
        driver = webdriver.Chrome(service=service, options=chrome_options)

        driver.get(web_address)
        time.sleep(3)

        total_width = driver.execute_script("return document.body.scrollWidth")
        total_height = driver.execute_script("return document.body.scrollHeight")
        driver.set_window_size(total_width, total_height)

        time.sleep(2)

        image_file = f"{recipe_name}.png"
        input_filename = f"{self.input_folder}/{image_file}"
        driver.save_screenshot(input_filename)
        driver.quit()
        time.sleep(2)

        self._convert_image_to_word(image_file)
        output_filename = f"{self.word_folder}/{recipe_name}.docx"
        return output_filename


def main():
    converter = RecipeConverter()
    converter.run()


if __name__ == "__main__":
    main()
